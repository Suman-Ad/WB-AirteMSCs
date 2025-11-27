# create_hsd_template.py
from docx import Document

doc = Document()

doc.add_heading('HSD RECEIVING FORMAT', level=1)
doc.add_paragraph('Document: HSD Receiving SOP').bold = True
doc.add_paragraph('Version No: Nxtra/Sec/V1.0')
doc.add_paragraph('Date of Release: 26th October 2022')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Site Name: ')
p.add_run('{{siteName}}')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Date: ')
p.add_run('{{date}}')

doc.add_paragraph()
doc.add_paragraph('Diesel Tanker')

# Core HSD Receiving Info table (keeps placeholders as single runs)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
rows = [
    ("Diesel Tanker (In time HH:MM)", "{{inTime}}"),
    ("Time: Informed to O&M team of arrival of HSD tanker by Security Team", "{{informTime}}"),
    ("Availability of Calibrated Dipstick brought by Tanker driver (Yes/No)", ""),
    ("HSD parking and ignition off time at HSD yard (HH:MM)", "{{parkingTime}}"),
    ("HSD Tanker settling time (HH:MM)", "{{fillingStartTime}}"),
    ("HSD Water availability check (Yes/No)", ""),
    ("Water quantity if available", ""),
    ("Flow meter/Dipstick reading before unloading in HSD tank", "")
]
for i, (left, right) in enumerate(rows):
    row = table.rows[i]
    row.cells[0].paragraphs[0].add_run(left)
    row.cells[1].paragraphs[0].add_run(right)

doc.add_paragraph()

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
rows2 = [
    ("Density", "{{density}}"),
    ("Temperature °C", "{{temperature}}"),
    ("Flow meter/Dip stick reading after unloading in HSD tank", ""),
    ("Actual HSD received (Difference)", "{{ltrs}}"),
    ("Accepted Invoice/Delivery challan number", "{{dillerInvoice}}")
]
for i, (left, right) in enumerate(rows2):
    r = table2.rows[i]
    r.cells[0].paragraphs[0].add_run(left)
    r.cells[1].paragraphs[0].add_run(right)

doc.add_paragraph()
doc.add_paragraph('HSD tanker (Out time) (HH:MM): ').add_run('{{outTime}}')

# Signature blocks
doc.add_paragraph()
doc.add_paragraph('O&M Team Name and Sign:')
doc.add_paragraph('Name: ').add_run('{{omName}}')
doc.add_paragraph('Sign: ').add_run('{{omSign}}')

doc.add_paragraph()
doc.add_paragraph('Security Name and Sign:')
doc.add_paragraph('Name: ').add_run('{{securityName}}')
doc.add_paragraph('Sign: ').add_run('{{securitySign}}')

doc.add_paragraph()
doc.add_paragraph('Checked by DC Manager')
doc.add_paragraph('Name: ').add_run('{{managerName}}')
doc.add_paragraph('Sign: ').add_run('{{managerSign}}')

# Save file
doc.save('HSD_Receiving_Template.docx')
print("Saved HSD_Receiving_Template.docx in current folder.")
